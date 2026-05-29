"""
generate_word.py
----------------
Genera BASE_MATEMATICA.docx — Guión completo de presentación del proyecto.
Ejecutar desde la raíz:
    /usr/bin/python3.13 generate_word.py
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ─────────────────────────────────────────────────────────────────────────────
# Utilidades de formato
# ─────────────────────────────────────────────────────────────────────────────

AZUL_OSCURO   = RGBColor(0x1F, 0x35, 0x64)
AZUL_MEDIO    = RGBColor(0x2E, 0x54, 0x96)
AZUL_CLARO    = RGBColor(0x2E, 0x74, 0xB5)
GRIS_TEXTO    = RGBColor(0x40, 0x40, 0x40)
BLANCO        = RGBColor(0xFF, 0xFF, 0xFF)

_TABLE_STYLE   = "Table Grid"
_LABEL_ARCHIVO = "Archivo: "
_SI_D0         = "Sí (D=0)"
_NO_D1         = "No (D=1)"
_COL_IDX       = "Índice"


def _set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def code_block(doc, code: str):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = _TABLE_STYLE
    cell = tbl.cell(0, 0)
    _set_cell_bg(cell, "F2F2F2")
    cell.paragraphs[0].clear()
    for i, line in enumerate(code.strip().split("\n")):
        p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        run = p.add_run(line)
        run.font.name = "Courier New"
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
    doc.add_paragraph()


def formula_box(doc, text: str):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = _TABLE_STYLE
    cell = tbl.cell(0, 0)
    _set_cell_bg(cell, "FFF9E6")
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.name = "Cambria Math"
    run.font.size = Pt(11)
    run.italic = True
    doc.add_paragraph()


def note_box(doc, text: str, bg="E8F4FD", label="NOTA"):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = _TABLE_STYLE
    cell = tbl.cell(0, 0)
    _set_cell_bg(cell, bg)
    p = cell.paragraphs[0]
    run_label = p.add_run(f"{label}: ")
    run_label.bold = True
    run_label.font.size = Pt(10)
    p.add_run(text).font.size = Pt(10)
    doc.add_paragraph()


def presentation_box(doc, text: str):
    """Caja naranja = 'cómo explicar esto en la presentación'."""
    note_box(doc, text, bg="FCE8D5", label="CÓMO EXPLICARLO")


def question_box(doc, q: str, a: str):
    """Caja verde = pregunta frecuente."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = _TABLE_STYLE
    cell = tbl.cell(0, 0)
    _set_cell_bg(cell, "E8F9E8")
    p = cell.paragraphs[0]
    r1 = p.add_run("PREGUNTA: ")
    r1.bold = True
    r1.font.size = Pt(10)
    p.add_run(q).font.size = Pt(10)
    p2 = cell.add_paragraph()
    r2 = p2.add_run("RESPUESTA: ")
    r2.bold = True
    r2.font.size = Pt(10)
    p2.add_run(a).font.size = Pt(10)
    doc.add_paragraph()


def warning_box(doc, text: str):
    note_box(doc, text, bg="FFE8E8", label="IMPORTANTE")


def add_table(doc, headers, rows):
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = _TABLE_STYLE
    for i, h in enumerate(headers):
        cell = tbl.cell(0, i)
        _set_cell_bg(cell, "4472C4")
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.color.rgb = BLANCO
        run.font.size = Pt(10)
    for r, row in enumerate(rows):
        bg = "FFFFFF" if r % 2 == 0 else "EEF3FB"
        for c, val in enumerate(row):
            cell = tbl.cell(r + 1, c)
            _set_cell_bg(cell, bg)
            cell.paragraphs[0].add_run(str(val)).font.size = Pt(9.5)
    doc.add_paragraph()


def h1(doc, text):
    p = doc.add_heading(text, level=1)
    p.runs[0].font.color.rgb = AZUL_OSCURO


def h2(doc, text):
    p = doc.add_heading(text, level=2)
    p.runs[0].font.color.rgb = AZUL_MEDIO


def h3(doc, text):
    p = doc.add_heading(text, level=3)
    p.runs[0].font.color.rgb = AZUL_CLARO


def body(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    return p


def bold_run(doc, label, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(label)
    r.bold = True
    r.font.size = Pt(11)
    p.add_run(text).font.size = Pt(11)


def bullet(doc, items: list):
    for item in items:
        p = doc.add_paragraph(item, style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
    doc.add_paragraph()


def numbered(doc, items: list):
    for item in items:
        p = doc.add_paragraph(item, style="List Number")
        p.paragraph_format.space_after = Pt(2)
    doc.add_paragraph()


def part_banner(doc, number: str, title: str):
    doc.add_page_break()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"PARTE {number}\n{title}")
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = AZUL_OSCURO
    doc.add_paragraph()


# ─────────────────────────────────────────────────────────────────────────────
# DOCUMENTO PRINCIPAL
# ─────────────────────────────────────────────────────────────────────────────

def build():
    doc = Document()
    for section in doc.sections:
        section.left_margin  = Cm(2.5)
        section.right_margin = Cm(2.5)
        section.top_margin   = Cm(2.5)
        section.bottom_margin = Cm(2.5)
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)

    # ─────────────────────────────────────────────────────────────
    # PORTADA
    # ─────────────────────────────────────────────────────────────
    for _ in range(3):
        doc.add_paragraph()

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("Bipartición Óptima en\nSistemas Causales Binarios")
    r.bold = True
    r.font.size = Pt(28)
    r.font.color.rgb = AZUL_OSCURO

    doc.add_paragraph()
    s = doc.add_paragraph()
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rs = s.add_run("Guión completo de presentación del proyecto\nAnálisis y Diseño de Algoritmos · 2025")
    rs.font.size = Pt(14)
    rs.font.color.rgb = GRIS_TEXTO

    doc.add_paragraph()
    note_box(doc,
        "Este documento es el guión de presentación. Cada sección explica: "
        "(1) el concepto matemático y de dónde viene, "
        "(2) cómo el código lo implementa exactamente, "
        "(3) cómo explicarlo en la presentación, y "
        "(4) posibles preguntas del jurado.",
        bg="E8F4FD", label="USO DEL DOCUMENTO"
    )

    doc.add_page_break()

    # ─────────────────────────────────────────────────────────────
    # PARTE I — EL PROBLEMA
    # ─────────────────────────────────────────────────────────────
    part_banner(doc, "I", "EL PROBLEMA: ¿QUÉ QUEREMOS RESOLVER?")

    # ── 1. La pregunta central ─────────────────────────────────
    h1(doc, "1. La Pregunta Central del Proyecto")

    body(doc,
        "Imagina un sistema con n partes (neuronas, genes, interruptores…) que cada una "
        "puede estar ENCENDIDA (1) o APAGADA (0). El sistema evoluciona en el tiempo: "
        "el estado actual determina probabilísticamente el estado siguiente."
    )
    body(doc,
        "La pregunta central es: ¿existe alguna forma de dividir este sistema en "
        "dos (o más) grupos independientes, de tal forma que lo que le pase al "
        "grupo 1 no afecte al grupo 2 y viceversa?"
    )

    formula_box(doc,
        "¿Existen S1, S2 tales que el sistema S1 ⊗ S2 se comporta igual que el sistema completo?"
    )

    presentation_box(doc,
        "Dile al jurado: 'Pensemos en el cerebro. Si puedo separar un grupo de neuronas "
        "que funcionan de manera completamente independiente al resto, entonces el sistema "
        "no está integrado. La bipartición óptima encuentra la mejor forma de hacer esa "
        "separación, y la discrepancia δ mide qué tan buena o mala es la separación.'"
    )

    question_box(doc,
        "¿Para qué sirve saber si el sistema es separable?",
        "Si δ = 0 el sistema es perfectamente separable: puede reducirse a dos subsistemas "
        "independientes sin perder información dinámica. Si δ > 0, hay dependencia causal "
        "irreducible entre las partes — concepto central en la teoría de la conciencia integrada (IIT)."
    )

    # ── 2. El pipeline completo ────────────────────────────────
    h1(doc, "2. El Pipeline Computacional (Resumen General)")

    body(doc,
        "El proyecto implementa los siguientes pasos en orden. Cada uno transforma "
        "la representación matemática del sistema:"
    )

    add_table(doc,
        ["Paso", "Nombre", "¿Qué hace?", "Módulo"],
        [
            ["1", "Cargar TPM", "Lee la matriz de probabilidades del CSV", "tpm_loader.py"],
            ["2", "Condicionar", "Fija las variables de fondo en su valor inicial", "conditioning.py"],
            ["3", "Marginalizar", "Reduce la TPM al subsistema candidato", "marginalization.py"],
            ["4", "Estado-nodo", "Descompone en distribuciones por variable", "state_node.py"],
            ["5", "Hipercubo / T", "Calcula la tabla de costos geométrica", "hypercube.py"],
            ["6", "EMD", "Mide discrepancia entre distribuciones", "emd.py"],
            ["7a", "Biparticiones", "Busca la mejor división en 2 grupos (k=2)", "bipartition.py"],
            ["7b", "K-particiones", "Busca la mejor división en k grupos (k≥2)", "kpartition.py"],
            ["7c", "GeoMIP", "Búsqueda geométrica eficiente usando T", "geometric.py"],
        ]
    )

    # ─────────────────────────────────────────────────────────────
    # PARTE II — EL SISTEMA Y SU REPRESENTACIÓN
    # ─────────────────────────────────────────────────────────────
    part_banner(doc, "II", "EL SISTEMA Y SU REPRESENTACIÓN")

    # ── 3. Variables binarias ──────────────────────────────────
    h1(doc, "3. Variables Binarias y el Espacio de Estados")

    h2(doc, "3.1 ¿Qué es una variable binaria?")
    body(doc,
        "Una variable binaria Xᵢ solo puede valer 0 o 1. En el proyecto representan "
        "cualquier proceso que tiene dos estados posibles:"
    )
    add_table(doc,
        ["Dominio", "Xᵢ = 0", "Xᵢ = 1"],
        [
            ["Neurociencia", "Neurona inactiva", "Neurona disparando"],
            ["Genómica", "Gen silenciado", "Gen expresado"],
            ["Lógica digital", "Bit apagado", "Bit encendido"],
            ["Sistemas físicos", "Estado bajo", "Estado alto"],
        ]
    )

    h2(doc, "3.2 El espacio de estados")
    body(doc,
        "Con n variables binarias hay exactamente 2ⁿ estados posibles. "
        "Cada estado es un vector binario de longitud n: [b₁, b₂, …, bₙ]."
    )

    add_table(doc,
        ["n", "Descripción", "2ⁿ estados"],
        [
            ["2", "4 estados (AB)", "00, 01, 10, 11"],
            ["3", "8 estados (ABC)", "000, 001, 010, ..., 111"],
            ["4", "16 estados (ABCD)", "0000, 0001, ..., 1111"],
            ["10", "1024 estados", "1,024 combinaciones posibles"],
            ["20", "1,048,576 estados", "≈ 1 millón"],
        ]
    )

    presentation_box(doc,
        "Explica: 'Pensemos en 3 interruptores (A, B, C). Cada uno puede estar encendido "
        "o apagado. Las 8 combinaciones posibles (000, 001, ..., 111) son los 8 estados "
        "del sistema. La TPM nos dice: si estamos en el estado ABC=100, "
        "¿cuál es la probabilidad de que el sistema vaya a cada uno de los otros 7 estados?'"
    )

    # ── 4. Codificación big-endian ─────────────────────────────
    h1(doc, "4. Codificación Big-Endian: De Bits a Índices")

    h2(doc, "4.1 El problema: ¿cómo numerar los estados?")
    body(doc,
        "Necesitamos asignar un número entero (índice) a cada estado binario para "
        "usarlo como índice de fila/columna en la matriz. La convención elegida es "
        "BIG-ENDIAN: el primer elemento del vector es el bit más significativo."
    )

    h2(doc, "4.2 La fórmula y de dónde viene")
    body(doc,
        "Esta es la misma fórmula que se usa para convertir cualquier número en "
        "base 2 a base 10. 'Big-endian' significa que el primer elemento (b₁) "
        "representa el bit de mayor peso (2^(n-1)):"
    )

    formula_box(doc,
        "estado [b₁, b₂, …, bₙ]  →  índice = b₁·2^(n-1) + b₂·2^(n-2) + … + bₙ·2⁰"
    )

    body(doc, "Ejemplos para n = 3 variables (A, B, C):")
    add_table(doc,
        ["Estado [A, B, C]", "Cálculo", _COL_IDX],
        [
            ["[0, 0, 0]", "0·4 + 0·2 + 0·1", "0"],
            ["[0, 0, 1]", "0·4 + 0·2 + 1·1", "1"],
            ["[0, 1, 0]", "0·4 + 1·2 + 0·1", "2"],
            ["[0, 1, 1]", "0·4 + 1·2 + 1·1", "3"],
            ["[1, 0, 0]", "1·4 + 0·2 + 0·1", "4"],
            ["[1, 0, 1]", "1·4 + 0·2 + 1·1", "5"],
            ["[1, 1, 0]", "1·4 + 1·2 + 0·1", "6"],
            ["[1, 1, 1]", "1·4 + 1·2 + 1·1", "7"],
        ]
    )

    h2(doc, "4.3 Implementación exacta en el código")
    bold_run(doc, _LABEL_ARCHIVO, "src/tpm_loader.py")
    code_block(doc, """\
def state_to_index(state: list[int]) -> int:
    # Acumulación bit a bit: desplaza 1 a la izquierda y agrega el siguiente bit
    result = 0
    for bit in state:
        result = (result << 1) | int(bit)
    return result
    # [1,0,0]: 0<<1|1=1, 1<<1|0=2, 2<<1|0=4 → índice 4 ✓

def index_to_state(idx: int, n: int) -> list[int]:
    # Extrae cada bit de derecha a izquierda usando desplazamiento y máscara
    return [(idx >> (n - 1 - i)) & 1 for i in range(n)]
    # idx=4, n=3: (4>>2)&1=1, (4>>1)&1=0, (4>>0)&1=0 → [1,0,0] ✓

# Caso especial importante: estado vacío → siempre índice 0
state_to_index([])   # → 0  (usado en biparticiones con parte sin variables en t)""")

    presentation_box(doc,
        "Cuando te pregunten por la codificación, di: 'Es exactamente la conversión "
        "de binario a decimal que aprendemos en fundamentos de programación. "
        "Big-endian significa que el primer elemento tiene mayor peso — igual que "
        "en el número 542: el 5 tiene peso 100, el 4 tiene peso 10.'"
    )

    question_box(doc,
        "¿Por qué big-endian y no little-endian?",
        "Es una convención arbitraria, pero big-endian es más natural para lectura humana "
        "(el primer dígito es el más significativo, igual que en los números decimales). "
        "Lo importante es ser consistente en todo el código, y así está definido en tpm_loader.py."
    )

    # ─────────────────────────────────────────────────────────────
    # PARTE III — LA TPM
    # ─────────────────────────────────────────────────────────────
    part_banner(doc, "III", "LA TPM: CORAZÓN DEL SISTEMA")

    # ── 5. Qué es y cómo se construye ─────────────────────────
    h1(doc, "5. La Matriz de Probabilidad de Transición (TPM)")

    h2(doc, "5.1 ¿Qué es la TPM y por qué la necesitamos?")
    body(doc,
        "La TPM (Transition Probability Matrix) captura la dinámica completa del sistema "
        "en una sola estructura. Es una matriz cuadrada P de tamaño 2ⁿ × 2ⁿ donde:"
    )

    formula_box(doc,
        "P[i][j]  =  P(V_{t+1} = j  |  V_t = i)\n\n"
        "La fila i describe: 'dado que el sistema está en el estado i ahora (tiempo t),\n"
        "¿cuál es la probabilidad de estar en cada estado j en el siguiente paso (t+1)?'"
    )

    body(doc,
        "La TPM es la fotografía completa de la dinámica del sistema. Con ella podemos:"
    )
    bullet(doc, [
        "Predecir la distribución de estados futuros",
        "Calcular cómo evoluciona el sistema desde cualquier estado inicial",
        "Medir la dependencia causal entre partes del sistema",
        "Encontrar la bipartición óptima",
    ])

    h2(doc, "5.2 Propiedades obligatorias")
    add_table(doc,
        ["Propiedad", "Fórmula", "¿Qué significa?"],
        [
            ["No negatividad", "P[i][j] ≥ 0", "Las probabilidades no pueden ser negativas"],
            ["Fila estocástica", "Σⱼ P[i][j] = 1", "Desde cualquier estado el sistema siempre va a algún lado"],
            ["Dimensión correcta", "2ⁿ × 2ⁿ", "El número de filas y columnas debe ser una potencia de 2"],
        ]
    )

    h2(doc, "5.3 Tipos de TPM")
    add_table(doc,
        ["Tipo", "Qué tiene en cada fila", "Ejemplo"],
        [
            ["Determinista", "Un solo 1 y el resto 0s", "Desde 000 siempre va a 010"],
            ["Estocástica", "Probabilidades distribuidas", "Desde 000: 60% va a 010, 40% va a 110"],
            ["Mixta", "Algunas filas deterministas, otras no", "Depende de cada estado"],
        ]
    )

    h2(doc, "5.4 Cómo se construye la TPM desde el mundo real")
    body(doc,
        "La TPM se puede construir de tres formas:"
    )
    numbered(doc, [
        "EMPÍRICA: Observando el sistema durante muchos pasos de tiempo y contando "
        "cuántas veces el estado i fue seguido por el estado j. TPM[i][j] = conteo(i→j) / conteo(i).",
        "ANALÍTICA: Derivando las reglas de actualización matemáticamente "
        "(por ejemplo, de ecuaciones booleanas o redes de Hopfield).",
        "DESDE ARCHIVO: En el proyecto, la TPM viene en un CSV donde cada fila "
        "i ya contiene las probabilidades P(→j) para j=0,1,...,2ⁿ-1.",
    ])

    h2(doc, "5.5 Cómo se lee el CSV y qué validaciones hace el código")
    bold_run(doc, _LABEL_ARCHIVO, "src/tpm_loader.py — función load_tpm")
    code_block(doc, """\
def load_tpm(filepath: str) -> tuple[np.ndarray, int]:
    tpm = np.loadtxt(filepath, delimiter=',')   # Lee el CSV sin encabezados
    n_states, n_cols = tpm.shape

    # Validación 1: debe ser cuadrada
    assert n_states == n_cols, "La TPM debe ser cuadrada (2^n × 2^n)"

    # Validación 2: el número de estados debe ser potencia de 2
    n = int(round(np.log2(n_states)))
    assert 2**n == n_states, f"Se necesita potencia de 2; hay {n_states} filas"

    # Validación 3: cada fila debe sumar 1 (distribución de probabilidad)
    assert np.allclose(tpm.sum(axis=1), 1.0, atol=1e-6), \
        "Cada fila debe sumar exactamente 1"

    return tpm, n   # retorna la matriz y el número de variables""")

    h2(doc, "5.6 Ejemplo concreto: sistema ABCD (data/tpm_abcd.csv)")
    body(doc,
        "Con 4 variables hay 2⁴ = 16 estados. La TPM es 16×16. "
        "Interpretando algunas filas (las que no son triviales):"
    )
    add_table(doc,
        ["Estado t (ABCD)", _COL_IDX, "Estado t+1 más probable", "P(→)", "Interpretación"],
        [
            ["0000", "0",  "0000", "1.0", "Estado absorbente — siempre permanece en 0000"],
            ["0001", "1",  "1000", "1.0", "D=1 activa A en el siguiente paso"],
            ["0010", "2",  "0010", "1.0", "Estado estable — C=1 se mantiene solo"],
            ["1000", "8",  "1101", "1.0", "A=1 activa A, C, D pero no B"],
            ["1001", "9",  "1101", "1.0", "A=1 domina; D sólo importa en t"],
        ]
    )

    presentation_box(doc,
        "Para presentar una fila de la TPM: 'La fila del estado 1000 dice: "
        "si A está encendida y todo lo demás apagado, en el siguiente instante "
        "el sistema estará en el estado 1101 con certeza absoluta (probabilidad 1). "
        "Eso significa que A activa a C y D, pero no activa a B.'"
    )

    # ── 6. Evolución del sistema ────────────────────────────────
    h1(doc, "6. Cómo Cambian las Probabilidades: Evolución del Sistema")

    h2(doc, "6.1 Un paso de tiempo")
    body(doc,
        "Si el sistema está con certeza en el estado s₀ en el tiempo t=0, "
        "entonces en t=1 está en el estado j con probabilidad P[s₀][j] "
        "(la fila s₀ de la TPM). En t=2 se distribuye según:"
    )
    formula_box(doc,
        "π(t) = π(0) · TPMᵗ\n\n"
        "donde π(0) es el vector fila con 1 en la posición del estado inicial y 0 en el resto.\n"
        "Cada multiplicación por TPM es un paso de tiempo."
    )

    h2(doc, "6.2 Ejemplo de evolución: ABCD desde estado 1000")
    code_block(doc, """\
# Estado inicial: 1000 (idx=8)
pi = np.zeros(16)
pi[8] = 1.0

for t in range(1, 4):
    pi = pi @ tpm          # multiplicar fila por matriz = un paso de tiempo
    # t=1: todo en estado 1101 (idx=13)  → P(1101)=1.0
    # t=2: todo en estado 0000 (idx=0)   → P(0000)=1.0  ← atractor
    # t=3: todo en estado 0000           → permanece

# El sistema cae al atractor 0000 en 2 pasos desde 1000""")

    question_box(doc,
        "¿Qué es un 'atractor' en el contexto de la TPM?",
        "Un estado o conjunto de estados al que el sistema converge inevitablemente. "
        "Se detecta como una fila i donde P[i][i] = 1 (el sistema permanece ahí para siempre), "
        "o un ciclo donde el sistema rota entre varios estados. "
        "En tpm_abcd.csv el estado 0000 es un atractor: P[0][0] = 1."
    )

    # ─────────────────────────────────────────────────────────────
    # PARTE IV — PREPARAR EL SISTEMA
    # ─────────────────────────────────────────────────────────────
    part_banner(doc, "IV", "PREPARAR EL SISTEMA: CONDICIONAMIENTO Y MARGINALIZACIÓN")

    # ── 7. Condicionamiento ────────────────────────────────────
    h1(doc, "7. Condicionamiento: Fijar las Variables de Fondo")

    h2(doc, "7.1 ¿Por qué condicionar?")
    body(doc,
        "No siempre queremos analizar el sistema completo V = {X₁,…,Xₙ}. "
        "Muchas veces nos interesa un SUBSISTEMA CANDIDATO Vᶜ ⊆ V "
        "(por ejemplo, sólo las variables ABC de un sistema ABCD)."
    )
    body(doc,
        "Las variables no estudiadas Vᵇ = V \\ Vᶜ se llaman 'variables de fondo' "
        "(background conditions). El condicionamiento las FIJA en su valor inicial, "
        "como si fueran constantes del experimento."
    )

    formula_box(doc,
        "TPM_cond = { fila i de TPM : ∀ k ∈ Vᵇ,  X_k(t) = valor inicial de k }"
    )

    h2(doc, "7.2 Ejemplo: sistema ABCD con estado inicial 1000, D es fondo")
    body(doc,
        "Estado inicial = [1,0,0,0] → el valor inicial de D (índice 3) es 0. "
        "Se seleccionan solo las filas donde el bit de D = 0:"
    )
    add_table(doc,
        ["Estado ABCD", _COL_IDX, "¿D=0?", "¿Se incluye?"],
        [
            ["0000", "0",  _SI_D0, "✅"],
            ["0001", "1",  _NO_D1, "❌"],
            ["0010", "2",  _SI_D0, "✅"],
            ["0011", "3",  _NO_D1, "❌"],
            ["...",  "...", "...",     "..."],
            ["1110", "14", _SI_D0, "✅"],
            ["1111", "15", _NO_D1, "❌"],
        ]
    )
    body(doc,
        "Resultado: 8 filas seleccionadas (los 8 estados con D=0). "
        "La TPM condicionada tiene forma 8 × 16."
    )

    h2(doc, "7.3 Implementación")
    bold_run(doc, _LABEL_ARCHIVO, "src/conditioning.py — función condition_tpm")
    code_block(doc, """\
for state_idx in range(2 ** n):
    state = index_to_state(state_idx, n)
    # ¿Todas las variables de fondo tienen su valor inicial?
    if all(state[bg] == bg_val
           for bg, bg_val in zip(background_vars, background_values)):
        selected_rows.append(state_idx)

conditioned_tpm = tpm[selected_rows, :]   # forma resultante: (2^n_cand, 2^n)""")

    # ── 8. Marginalización ─────────────────────────────────────
    h1(doc, "8. Marginalización: Reducir al Subsistema Candidato")

    h2(doc, "8.1 El problema que resuelve")
    body(doc,
        "Después del condicionamiento, la TPM tiene el número correcto de FILAS "
        "(estados del candidato Vᶜ en t) pero todavía tiene 2^n COLUMNAS "
        "(estados del sistema completo en t+1, incluyendo las variables de fondo)."
    )
    body(doc,
        "La marginalización en COLUMNAS elimina las variables de fondo de t+1, "
        "produciendo una TPM cuadrada del subsistema candidato: 2^n_c × 2^n_c."
    )

    h2(doc, "8.2 Marginalización de columnas (t+1): sumar")
    formula_box(doc,
        "TPM_marg[fila, j'] = Σ_{j : proj_{Vᶜ}(j) = j'}  TPM[fila, j]\n\n"
        "Se suman las columnas cuyos estados proyectan al mismo estado candidato j'.\n"
        "NO se divide: la suma ya es la probabilidad marginal correcta."
    )

    h2(doc, "8.3 Marginalización de filas (t): promediar")
    body(doc,
        "Al evaluar biparticiones, también necesitamos marginalizar FILAS "
        "(eliminar variables de t). Aquí se PROMEDIA (no suma) para preservar "
        "la propiedad de fila estocástica:"
    )
    formula_box(doc,
        "TPM_marg[i', j] = (1/|grupo_{i'}|) · Σ_{i : proj(i) = i'}  TPM[i, j]\n\n"
        "El promedio garantiza que cada fila del resultado siga sumando 1."
    )

    question_box(doc,
        "¿Por qué se suma en columnas pero se promedia en filas?",
        "Porque columnas y filas tienen interpretaciones distintas. "
        "Las columnas representan eventos del futuro (t+1): la probabilidad de ir al "
        "estado proyectado j' es la suma de todas las probabilidades de estados que "
        "proyectan a j'. "
        "Las filas representan condiciones del presente (t): si el estado presente "
        "puede ser i₁ ó i₂ (con igual frecuencia), la distribución condicional es el "
        "promedio de las distribuciones de i₁ e i₂."
    )

    h2(doc, "8.4 Implementación")
    bold_run(doc, _LABEL_ARCHIVO, "src/marginalization.py")
    code_block(doc, """\
# Marginalización de COLUMNAS (t+1) — sumar columnas con misma proyección
marginalized = np.zeros((n_rows, 2**n_keep))
for col_idx in range(2 ** n_total):
    full_state = index_to_state(col_idx, n_total)
    projected  = [full_state[v] for v in keep_vars]
    out_col    = state_to_index(projected)
    marginalized[:, out_col] += tpm[:, col_idx]    # SUMA

# Marginalización de FILAS (t) — promediar filas con misma proyección
accumulated = np.zeros((2**n_keep, n_cols))
counts      = np.zeros(2**n_keep)
for row_idx in range(n_rows):
    full_state = index_to_state(row_idx, n_candidate)
    projected  = [full_state[v] for v in keep_vars]
    out_row    = state_to_index(projected)
    accumulated[out_row] += tpm[row_idx]
    counts[out_row] += 1
result = accumulated / counts[:, None]              # PROMEDIO""")

    # ─────────────────────────────────────────────────────────────
    # PARTE V — INDEPENDENCIA Y RECONSTRUCCIÓN
    # ─────────────────────────────────────────────────────────────
    part_banner(doc, "V", "INDEPENDENCIA CONDICIONAL Y PRODUCTO TENSORIAL")

    # ── 9. Estado-nodo e independencia ─────────────────────────
    h1(doc, "9. Representación Estado-Nodo e Independencia Condicional")

    h2(doc, "9.1 El teorema clave")
    body(doc,
        "El teorema de independencia condicional establece que, conocido el estado "
        "completo V_t en t, las variables individuales en t+1 son INDEPENDIENTES entre sí. "
        "Este teorema es válido porque cada variable en t+1 depende de V_t de forma "
        "determinista (o probabilística) y ese V_t es la única fuente de dependencia:"
    )
    formula_box(doc,
        "P(X₁_{t+1}, X₂_{t+1}, …, Xₙ_{t+1} | V_t = s)\n\n"
        "= P(X₁_{t+1}|V_t=s) · P(X₂_{t+1}|V_t=s) · … · P(Xₙ_{t+1}|V_t=s)"
    )
    body(doc,
        "En otras palabras: condicionado en el estado actual completo, el futuro de "
        "cada variable es independiente del futuro de las demás."
    )

    h2(doc, "9.2 Las matrices estado-nodo")
    body(doc,
        "Gracias a este teorema, en lugar de guardar la distribución conjunta "
        "(una fila de 2ⁿ valores), alcanza con guardar n distribuciones marginales, "
        "una por variable. Cada matriz estado-nodo Mᵢ tiene forma (2ⁿ × 2):"
    )
    formula_box(doc,
        "Mᵢ[s, v] = P(Xᵢ_{t+1} = v | V_t = s),   v ∈ {0, 1}\n\n"
        "Columna 0: probabilidad de que Xᵢ valga 0 en t+1\n"
        "Columna 1: probabilidad de que Xᵢ valga 1 en t+1"
    )

    h2(doc, "9.3 Ejemplo: variable A del subsistema ABC")
    add_table(doc,
        ["Estado t (ABC)", "P(A=0 | t)", "P(A=1 | t)", "Interpretación"],
        [
            ["000", "1.000", "0.000", "A siempre será 0 si partimos de 000"],
            ["001", "0.000", "1.000", "A siempre será 1 si partimos de 001"],
            ["100", "0.000", "1.000", "A siempre será 1 si partimos de 100"],
            ["101", "0.500", "0.500", "A es incierta desde 101 (promedio de dos transiciones)"],
        ]
    )

    note_box(doc,
        "¿Por qué P(A=1|101) = 0.5? Porque el estado 101 en el candidato ABC "
        "corresponde a ABCD = 1010 y 1011 en el sistema completo. Estos dos estados "
        "tienen transiciones distintas respecto a A. Al marginalizar D, se promedian "
        "→ resultado 0.5. El código lo calcula sumando P(col) donde col_state[0]=1.",
        bg="FFF3CD", label="DETALLE TÉCNICO"
    )

    h2(doc, "9.4 Implementación")
    bold_run(doc, _LABEL_ARCHIVO, "src/state_node.py — state_state_to_state_node")
    code_block(doc, """\
def state_state_to_state_node(tpm, n):
    node_matrices = []
    for var_idx in range(n):
        node_mat = np.zeros((2**n, 2))
        for row in range(2**n):
            for col in range(2**n):
                col_state = index_to_state(col, n)
                xi_val = col_state[var_idx]         # ¿vale 0 o 1 la variable var_idx?
                node_mat[row, xi_val] += tpm[row, col]   # acumular
        node_matrices.append(node_mat)   # node_mat[s, 0] + node_mat[s, 1] = 1 ∀s
    return node_matrices""")

    # ── 10. Producto tensorial ─────────────────────────────────
    h1(doc, "10. Producto Tensorial: Reconstruir la Distribución Conjunta")

    h2(doc, "10.1 Del marginal al conjunto: producto de Kronecker")
    body(doc,
        "Si tenemos la distribución marginal de cada variable Xᵢ dado el estado s, "
        "podemos reconstruir la distribución conjunta usando el producto de Kronecker ⊗:"
    )
    formula_box(doc,
        "P(V_{t+1} | V_t = s)  =  M₁[s, :] ⊗ M₂[s, :] ⊗ … ⊗ Mₙ[s, :]\n\n"
        "Para n=2: [pA0, pA1] ⊗ [pB0, pB1]\n"
        "= [pA0·pB0,  pA0·pB1,  pA1·pB0,  pA1·pB1]\n"
        "= [P(AB=00), P(AB=01), P(AB=10), P(AB=11)]"
    )

    body(doc,
        "Esta reconstrucción es EXACTA (sin pérdida de información) gracias al teorema "
        "de independencia condicional. Si el sistema era determinista, la reconstrucción "
        "devuelve exactamente la TPM original."
    )

    h2(doc, "10.2 Producto tensorial para biparticiones: el núcleo del proyecto")
    body(doc,
        "Al evaluar una bipartición {S1, S2}, la pregunta es: "
        "'si asumimos que S1 y S2 son INDEPENDIENTES, ¿cuál sería la TPM resultante?' "
        "Esa TPM 'como si fueran independientes' se llama TPM RECONSTRUIDA:"
    )
    formula_box(doc,
        "P_rec(j | i) = P_S1(proj_{S1,t+1}(j) | proj_{S1,t}(i))\n"
        "             × P_S2(proj_{S2,t+1}(j) | proj_{S2,t}(i))\n\n"
        "P_S1 = TPM marginalizada a las variables de S1 en t y t+1\n"
        "P_S2 = idem para S2\n"
        "proj_{S,t}(i) = estado i proyectado a las variables de S en tiempo t"
    )

    h2(doc, "10.3 Implementación")
    bold_run(doc, _LABEL_ARCHIVO, "src/state_node.py — tensor_product_tpm")
    code_block(doc, """\
def tensor_product_tpm(tpm_full, n, s1_t, s1_t1, s2_t, s2_t1):
    # Marginalizar cada parte a sus propias variables
    tpm_s1 = marginalize_rows(tpm_full, n, s1_t)       # filas: vars de S1 en t
    tpm_s1 = marginalize_columns(tpm_s1, n, s1_t1)     # cols:  vars de S1 en t+1
    tpm_s2 = marginalize_rows(tpm_full, n, s2_t)
    tpm_s2 = marginalize_columns(tpm_s2, n, s2_t1)

    tpm_rec = np.zeros((2**n, 2**n))
    for i in range(2**n):
        fi = index_to_state(i, n)
        row_s1 = state_to_index([fi[v] for v in s1_t])   # proyectar i a S1
        row_s2 = state_to_index([fi[v] for v in s2_t])   # proyectar i a S2
        for j in range(2**n):
            fj = index_to_state(j, n)
            col_s1 = state_to_index([fj[v] for v in s1_t1])
            col_s2 = state_to_index([fj[v] for v in s2_t1])
            tpm_rec[i, j] = tpm_s1[row_s1, col_s1] * tpm_s2[row_s2, col_s2]
    return tpm_rec""")

    presentation_box(doc,
        "Para la presentación: 'La TPM reconstruida nos dice: si S1 y S2 fueran "
        "completamente independientes, ¿cómo sería la probabilidad de transición? "
        "La construimos multiplicando las probabilidades de cada parte por separado. "
        "Si la reconstruida es idéntica a la original, el sistema realmente es separable.'"
    )

    # ─────────────────────────────────────────────────────────────
    # PARTE VI — LA MÉTRICA: HAMMING, T Y EMD
    # ─────────────────────────────────────────────────────────────
    part_banner(doc, "VI", "LA MÉTRICA: HIPERCUBO, TABLA T Y EMD")

    # ── 11. Hipercubo y Hamming ────────────────────────────────
    h1(doc, "11. El Hipercubo n-dimensional y la Distancia de Hamming")

    h2(doc, "11.1 Representación geométrica del espacio de estados")
    body(doc,
        "Los 2ⁿ estados del sistema se pueden visualizar como los vértices de un "
        "HIPERCUBO n-dimensional. Cada dimensión del cubo corresponde a una variable. "
        "Dos estados son VECINOS (adyacentes) en el hipercubo si y sólo si "
        "difieren en exactamente 1 bit."
    )
    add_table(doc,
        ["n", "Figura geométrica", "Vértices", "Aristas"],
        [
            ["1", "Segmento", "2", "1"],
            ["2", "Cuadrado", "4", "4"],
            ["3", "Cubo", "8", "12"],
            ["4", "Hipercubo 4D (teseracto)", "16", "32"],
            ["n", "Hipercubo n-D", "2ⁿ", "n·2^(n-1)"],
        ]
    )

    h2(doc, "11.2 Distancia de Hamming")
    body(doc,
        "La distancia de Hamming entre dos estados es el número de bits en que difieren. "
        "En el hipercubo, equivale exactamente a la longitud del camino más corto."
    )
    formula_box(doc,
        "d_H(x, y) = |{i : xᵢ ≠ yᵢ}| = número de posiciones diferentes\n\n"
        "Implementación eficiente: d_H(i, j) = popcount(i XOR j)\n"
        "= bin(i ^ j).count('1')"
    )
    add_table(doc,
        ["Estado i (bin)", "Estado j (bin)", "i XOR j", "Bits diferentes", "d_H"],
        [
            ["000 (0)", "001 (1)", "001", "1 (sólo C)", "1"],
            ["000 (0)", "011 (3)", "011", "2 (B y C)", "2"],
            ["000 (0)", "111 (7)", "111", "3 (A, B, C)", "3"],
            ["101 (5)", "110 (6)", "011", "2 (B y C)", "2"],
        ]
    )

    # ── 12. Función de costo T ─────────────────────────────────
    h1(doc, "12. Función de Costo Geométrica T (GeoMIP Ec. 3.1)")

    h2(doc, "12.1 ¿Qué mide T?")
    body(doc,
        "Para cada variable Xᵢ, se asocia a cada vértice v del hipercubo el valor "
        "X[v] = P(Xᵢ_{t+1}=1 | V_t=v). Luego se define el costo de transición t(i,j) "
        "como la energía geométrica necesaria para 'mover' la probabilidad del "
        "vértice i al vértice j, considerando la topología del hipercubo."
    )
    body(doc,
        "A diferencia de la distancia de Hamming (puramente geométrica), "
        "la función T incorpora la información probabilística del sistema."
    )

    h2(doc, "12.2 La fórmula recursiva")
    formula_box(doc,
        "t(i, j) = γ · ( |X[i] − X[j]| + Σ_{k ∈ N_opt(i,j)} t(i, k) )\n\n"
        "γ = 2^(−d_H(i,j))               factor de decaimiento exponencial\n"
        "N_opt(i,j) = vecinos de i en caminos óptimos hacia j\n"
        "           = {k ∈ adj(i) : d_H(k, j) = d_H(i, j) − 1}"
    )
    body(doc,
        "Se calcula BOTTOM-UP: primero los pares con d_H=1, luego d_H=2, etc. "
        "Esto garantiza que t(i,k) siempre esté disponible cuando se calcula t(i,j)."
    )

    h2(doc, "12.3 Ejemplo de cálculo manual (n=2)")
    body(doc, "X = [0.0, 1.0, 0.0, 1.0] → desde el estado 00 (i=0):")
    add_table(doc,
        ["Par (i→j)", "d_H", "γ=2^(-d)", "|X[i]−X[j]|", "N_opt", "t(i,j)"],
        [
            ["00→01", "1", "0.50", "|0−1|=1", "—", "0.50·1 = 0.500"],
            ["00→10", "1", "0.50", "|0−0|=0", "—", "0.50·0 = 0.000"],
            ["00→11", "2", "0.25", "|0−1|=1", "{01, 10}", "0.25·(1+0.5+0) = 0.375"],
        ]
    )

    h2(doc, "12.4 Implementación del BFS bottom-up")
    bold_run(doc, _LABEL_ARCHIVO, "src/hypercube.py — compute_cost_table")
    code_block(doc, """\
for i in range(num_states):
    # d=1: sin recursión, sólo diferencia directa entre vértices vecinos
    for j in adj[i]:
        T[i, j] = 0.5 * abs(node_probs[i] - node_probs[j])

    # d=2..n: acumular vecinos ya calculados (garantizado por el orden BFS)
    for d in range(2, n + 1):
        gamma = 2.0 ** (-d)
        for j in range(num_states):
            if hamming_distance(i, j, n) != d:
                continue
            # Vecinos de i que están un paso más cerca de j
            neighbors_toward = [k for k in adj[i]
                                 if hamming_distance(k, j, n) == d - 1]
            neighbor_sum = sum(T[i, k] for k in neighbors_toward)
            T[i, j] = gamma * (abs(node_probs[i] - node_probs[j]) + neighbor_sum)""")

    # ── 13. EMD ────────────────────────────────────────────────
    h1(doc, "13. Earth Mover's Distance (EMD): Medir la Discrepancia")

    h2(doc, "13.1 La analogía de la arena")
    body(doc,
        "Imagina que P y Q son dos montones de arena distribuidos sobre los vértices "
        "del hipercubo. La EMD es el mínimo trabajo necesario para transformar "
        "la distribución P en la distribución Q, donde mover 1 unidad de arena "
        "por 1 arista del hipercubo (d_H=1) cuesta 1 unidad de trabajo."
    )

    h2(doc, "13.2 Formulación como Programación Lineal")
    formula_box(doc,
        "minimizar      Σᵢⱼ F[i,j] · d_H(i,j)\n\n"
        "sujeto a:      Σⱼ F[i,j] = P[i]     ∀ i    (conservar oferta)\n"
        "               Σᵢ F[i,j] = Q[j]     ∀ j    (satisfacer demanda)\n"
        "               F[i,j] ≥ 0\n\n"
        "Variables de decisión: F[i,j] = cantidad de masa movida de i a j\n"
        "Tamaño del problema: 4ⁿ variables, 2·2ⁿ restricciones"
    )

    h2(doc, "13.3 Implementación")
    bold_run(doc, _LABEL_ARCHIVO, "src/emd.py — función emd")
    code_block(doc, """\
from scipy.optimize import linprog

def emd(p, q, cost_matrix):
    n_src, n_dst = len(p), len(q)
    n_vars = n_src * n_dst
    c = cost_matrix.flatten()          # coeficientes de la función objetivo

    A_eq = np.zeros((n_src + n_dst, n_vars))
    for i in range(n_src):            # restricciones de oferta
        A_eq[i, i*n_dst:(i+1)*n_dst] = 1.0
    for j in range(n_dst):            # restricciones de demanda
        A_eq[n_src+j, j::n_dst] = 1.0

    b_eq = np.concatenate([p, q])
    result = linprog(c, A_eq=A_eq, b_eq=b_eq,
                     bounds=[(0,None)]*n_vars,
                     method='highs')   # HiGHS: solver LP de alta eficiencia
    return float(result.fun)""")

    question_box(doc,
        "¿Por qué usar EMD y no simplemente la diferencia L1 (|P-Q|)?",
        "La diferencia L1 no considera la estructura geométrica del espacio de estados. "
        "La EMD con distancia de Hamming sí: mover masa entre estados cercanos "
        "(d_H=1) cuesta menos que entre estados lejanos (d_H=n). "
        "Esto captura mejor la 'similitud causal' entre distribuciones."
    )

    # ─────────────────────────────────────────────────────────────
    # PARTE VII — BIPARTICIÓN Y K-PARTICIÓN ÓPTIMA
    # ─────────────────────────────────────────────────────────────
    part_banner(doc, "VII", "BIPARTICIÓN Y K-PARTICIÓN ÓPTIMA")

    # ── 14. Bipartición óptima ─────────────────────────────────
    h1(doc, "14. Bipartición Óptima: El Resultado Central")

    h2(doc, "14.1 Definición formal")
    body(doc,
        "Una bipartición divide el conjunto de pares (variable, tiempo) en dos partes. "
        "El conjunto total tiene 2n elementos: n variables en t y n variables en t+1."
    )
    formula_box(doc,
        "S1 ∪ S2 = {(X₁,t), …, (Xₙ,t), (X₁,t+1), …, (Xₙ,t+1)}\n"
        "S1 ∩ S2 = ∅,   S1 ≠ ∅,   S2 ≠ ∅\n\n"
        "Discrepancia:  δ({S1,S2}, s₀) = EMD( P(V_{t+1}|V_t=s₀),  P_rec(V_{t+1}|V_t=s₀) )\n\n"
        "Bipartición óptima:  {S1*, S2*} = argmin_{S1,S2} δ({S1,S2}, s₀)"
    )

    h2(doc, "14.2 El espacio de biparticiones y por qué congela el PC")
    body(doc,
        "Con 2n elementos en el conjunto, el número de biparticiones no triviales es:"
    )
    formula_box(doc,
        "|biparticiones| = 2^(2n−1) − 1\n\n"
        "Derivación: hay 2^(2n) subconjuntos totales. Descartando el vacío y el total: "
        "2^(2n)−2. Dividiendo por 2 para eliminar S1|S2 = S2|S1: (2^(2n)−2)/2 = 2^(2n−1)−1."
    )
    add_table(doc,
        ["n", "Biparticiones 2^(2n-1)-1", "EMDs a resolver", "Viabilidad"],
        [
            ["2", "7", "7", "Instantáneo"],
            ["3", "31", "31", "< 1 segundo"],
            ["5", "511", "511", "Segundos"],
            ["10", "524,287", "524,287", "Minutos a horas"],
            ["15", "536,870,911", "≈ 537 millones", "Días"],
            ["20", "549,755,813,887", "≈ 550 mil millones", "CONGELA EL PC"],
        ]
    )
    warning_box(doc,
        "Cada EMD resuelve un problema de programación lineal con 4ⁿ variables. "
        "Para n=20: 4²⁰ = 1.099 billones de variables por LP × 550 mil millones de LPs. "
        "El código tiene un límite MAX_BIP = 2000 para evitar este problema."
    )

    h2(doc, "14.3 Implementación de la búsqueda exhaustiva")
    bold_run(doc, _LABEL_ARCHIVO, "src/bipartition.py — find_optimal_bipartition")
    code_block(doc, """\
all_var_time = build_bipartition_variables(n)
# → [(0,'t'),(1,'t'),...,(n-1,'t'),(0,'t+1'),...,(n-1,'t+1')]

for s1, s2 in generate_bipartitions(all_var_time):
    if len(all_results) >= max_bipartitions:
        truncated = True; break

    s1_t, s1_t1, s2_t, s2_t1 = split_bipartition(s1, s2)

    # Construir TPM reconstruida suponiendo independencia entre S1 y S2
    tpm_rec = tensor_product_tpm(tpm, n, s1_t, s1_t1, s2_t, s2_t1)

    # Calcular discrepancia EMD entre distribución original y reconstruida
    p = tpm[initial_state_idx, :]    # distribución original
    q = tpm_rec[initial_state_idx, :]  # distribución reconstruida
    delta = emd(p, q, hamming_distance_matrix(n))

    if delta < best_delta:
        best_delta = delta
        best_s1, best_s2 = s1_label, s2_label""")

    presentation_box(doc,
        "Para explicar la búsqueda exhaustiva: 'Para cada forma posible de dividir "
        "el sistema en dos partes, calculamos: si estas dos partes fueran independientes, "
        "¿cómo se vería la dinámica? Luego comparamos esa dinámica hipotética con la "
        "real usando EMD. La división que produce la diferencia más pequeña es la "
        "bipartición óptima.'"
    )

    # ── 15. K-Particiones ─────────────────────────────────────
    h1(doc, "15. K-Particiones: Generalizar a Más de Dos Grupos")

    h2(doc, "15.1 ¿Qué es una k-partición?")
    body(doc,
        "Una k-partición divide el sistema en k partes disjuntas y no vacías "
        "(en lugar de solo 2). La bipartición es el caso especial k=2."
    )
    formula_box(doc,
        "P_rec(j | i) = ∏_{p=1}^{k}  P_p(proj_{p,t+1}(j) | proj_{p,t}(i))\n\n"
        "El producto de k factores (uno por parte) reemplaza el producto de 2 factores."
    )

    h2(doc, "15.2 Conteo: números de Stirling y de Bell")
    body(doc,
        "El número de k-particiones de m elementos es el número de Stirling de segunda "
        "especie S(m, k). El total de todas las k es el número de Bell B(m):"
    )
    add_table(doc,
        ["m=2n", "k=2 (bipart.)", "k=3", "k=4", "Bell B(m) (total)"],
        [
            ["4 (n=2)",  "7",       "6",       "1",       "15"],
            ["6 (n=3)",  "31",      "90",      "65",      "203"],
            ["8 (n=4)",  "127",     "966",     "4,060",   "4,140 (sin k=1)"],
            ["10 (n=5)", "511",     "22,827",  "357,423", "115,975"],
        ]
    )
    body(doc,
        "Para n=3 el espacio total es 203 particiones (vs 31 para biparticiones solo). "
        "Para n=5 ya hay 115,975 — y esto crece mucho más rápido que las biparticiones."
    )

    h2(doc, "15.3 Módulo kpartition.py — enumeración canónica")
    body(doc,
        "La enumeración evita duplicados simétricos usando FORMA CANÓNICA: "
        "el primer elemento siempre va al grupo 0, y los nuevos grupos aparecen "
        "en orden estrictamente creciente."
    )
    bold_run(doc, _LABEL_ARCHIVO, "src/kpartition.py")
    code_block(doc, """\
def generate_kpartitions(variables: list, k: int):
    \"\"\"Genera todas las k-particiones canónicas. Total = S(len(variables), k).\"\"\"
    # El primer elemento siempre va al grupo 0 (evita S1|S2 = S2|S1)
    yield from _generate_kpart_helper(list(variables), k, assignment=[0], n_used=1)

def find_optimal_kpartition(tpm, n, k, initial_state_idx, variable_names, ...):
    for partition in generate_kpartitions(all_var_time, k):
        parts_split = split_kpartition(partition)   # [(part_t, part_t1), ...]
        # TPM reconstruida como PRODUCTO de k factores (generaliza el caso k=2)
        tpm_rec = reconstruct_kpartition_tpm(tpm, n, parts_split)
        delta = emd(tpm[s0,:], tpm_rec[s0,:], D)
        if delta < best: best = delta; best_part = partition

# Conteo:
count_kpartitions(m=6, k=2)   # → 31  (S(6,2))
count_kpartitions(m=6, k=3)   # → 90  (S(6,3))
bell_number(m=6)               # → 203 (B(6))""")

    # ─────────────────────────────────────────────────────────────
    # PARTE VIII — SOLUCIÓN GEOMÉTRICA (GEOMIP)
    # ─────────────────────────────────────────────────────────────
    part_banner(doc, "VIII", "SOLUCIÓN GEOMÉTRICA: EL ALGORITMO GEOMIP")

    # ── 16. La idea geométrica ─────────────────────────────────
    h1(doc, "16. La Idea Detrás de la Solución Geométrica")

    h2(doc, "16.1 El problema de la fuerza bruta")
    body(doc,
        "La búsqueda exhaustiva evalúa TODAS las biparticiones/k-particiones, "
        "cada una con un LP (EMD). Para n grande esto es computacionalmente imposible."
    )
    body(doc,
        "La pregunta del GeoMIP es: ¿podemos usar la estructura de la tabla T para "
        "IDENTIFICAR DIRECTAMENTE cuáles biparticiones son buenas candidatas, "
        "sin evaluar todas?"
    )

    h2(doc, "16.2 La clave: patrones de costo cero en T")
    body(doc,
        "Para cada variable v, la fila T_v[s₀, :] contiene los costos de transición "
        "desde el estado inicial s₀ a todos los demás estados. Los estados j donde "
        "T_v[s₀, j] = 0 son 'estados equivalentes' para la variable v: "
        "la variable no distingue entre s₀ y j."
    )
    body(doc,
        "La idea central del GeoMIP (Sección 4.2.4): "
        "variables que comparten el mismo patrón de 'estados cero' desde s₀ "
        "deben ir en el MISMO grupo de la partición. "
        "Variables con patrones distintos van en grupos distintos."
    )

    note_box(doc,
        "¿Cuándo T_v[s₀, j] = 0? Para d_H=1: cuando X_v[s₀] = X_v[j], es decir, "
        "P(Xv=1|s₀) = P(Xv=1|j). Para d>1: cuando la diferencia se propaga a cero "
        "a través de todos los vecinos intermedios. "
        "En la práctica: estados j que son 'indistinguibles' para la variable v desde s₀.",
        bg="FFF3CD", label="DETALLE TÉCNICO"
    )

    # ── 17. El algoritmo GeoMIP ────────────────────────────────
    h1(doc, "17. El Algoritmo GeoMIP (Algorithm 2)")

    body(doc, "El algoritmo completo en pseudocódigo (GeoMIP Capítulo 5):")
    code_block(doc, """\
Algorithm 2 — Algoritmo Geométrico
─────────────────────────────────────────────────────────────────
Input:  Subsistema S con n variables, estado inicial s₀
Output: Partición óptima P_opt

1.  tensors ← DescomponerEnTensores(S)
    → Matrices estado-nodo para cada variable (ya en state_node.py)

2.  T ← InicializarTablaDeTransiciones()
    for cada variable v:
        for cada par de estados (i, j):
            T[v, i, j] ← CalcularCostoDeTransicion(i, j, tensors[v])
    → Ya implementado en hypercube.py::compute_cost_table

3.  candidates ← IdentificarParticionesCandidatas(T, s₀, k)
    → Agrupar variables por su patrón de ceros en T[:,s₀,:]
    → Generar variantes intercambiando t/t+1

4.  P_opt ← EvaluarCandidatos(candidates)
    → Calcular EMD exacto para cada candidato
    → Retornar el de menor δ
─────────────────────────────────────────────────────────────────
Complejidad: O(n·2ⁿ) para construir T  vs  O(2^(2n-1)) exhaustivo""")

    h2(doc, "17.1 Implementación de los candidatos geométricos")
    bold_run(doc, _LABEL_ARCHIVO, "src/geometric.py — identify_geometric_candidates")
    code_block(doc, """\
def identify_geometric_candidates(tpm, n, initial_state_idx, k, cost_tables):
    # Para cada variable v, encontrar estados j con T_v[s0, j] ≈ 0
    patterns = {}
    for v in range(n):
        row = cost_tables[v][initial_state_idx, :].copy()
        row[initial_state_idx] = np.inf           # ignorar diagonal (trivial)
        zeros = frozenset(j for j in range(2**n) if abs(row[j]) < 1e-9)
        if not zeros:
            zeros = frozenset([int(np.argmin(row))])   # mínimo si no hay ceros exactos
        patterns[zeros].append(v)

    # Agrupar variables por su patrón de ceros
    unique_groups = list(patterns.values())
    # → variables con el mismo patrón = mismo grupo de la partición

    # Ajustar al número k pedido (fusionar o dividir grupos)
    while len(unique_groups) < k:
        largest = max(unique_groups, key=len)
        mid = len(largest) // 2
        unique_groups.remove(largest)
        unique_groups.extend([largest[:mid], largest[mid:]])

    while len(unique_groups) > k:
        a = min(unique_groups, key=len)
        b = min((g for g in unique_groups if g is not a), key=len)
        unique_groups = [g for g in unique_groups if g is not a and g is not b]
        unique_groups.append(a + b)

    # Construir candidato y variantes (intercambiar t/t+1 de cada variable)
    return [candidato_base] + variantes""")

    h2(doc, "17.2 La clase GeometricSIA")
    body(doc,
        "Como requieren los entregables del GeoMIP (Capítulo 5), se implementó "
        "la clase GeometricSIA con la interfaz estándar:"
    )
    bold_run(doc, _LABEL_ARCHIVO, "src/controllers/strategies/geometric.py")
    code_block(doc, """\
class GeometricSIA:
    def __init__(self, tpm, n, initial_state_idx, variable_names):
        self.tpm = tpm
        self.n   = n
        self.initial_state_idx = initial_state_idx
        # Paso 1: Descomponer en tensores elementales
        self.node_matrices = state_state_to_state_node(tpm, n)
        # Paso 2: Calcular T para todas las variables O(n·2^n)
        self.cost_tables = _compute_all_cost_tables(tpm, n)

    def aplicar_estrategia(self, k=2) -> dict:
        \"\"\"Pasos 3 y 4 del Algorithm 2.\"\"\"
        candidates = self.identificar_candidatos(k)         # Paso 3
        # Paso 4: Evaluar con EMD exacto
        scored = [(label, emd(p_orig, p_rec, D)) for cand in candidates ...]
        best_label, best_delta = min(scored, key=lambda x: x[1])
        return {'k': k, 'optimal_partition': best_label, 'min_delta': best_delta, ...}

    def calcular_transicion_coste(self, v, i, j) -> float:
        return float(self.cost_tables[v][i, j])    # T_v[i, j]

    def identificar_candidatos(self, k=2) -> list:
        return identify_geometric_candidates(self.tpm, self.n,
                                              self.initial_state_idx, k,
                                              self.cost_tables)""")

    # ── 18. Comparación ───────────────────────────────────────
    h1(doc, "18. Comparación: Exhaustivo vs Geométrico")

    add_table(doc,
        ["Aspecto", "Búsqueda Exhaustiva", "Solución Geométrica (GeoMIP)"],
        [
            ["Biparticiones evaluadas", "2^(2n-1)-1  (todas)", "~n·k candidatos (muy pocas)"],
            ["Costo por evaluación", "1 LP completo (EMD)", "1 LP completo (EMD)"],
            ["Complejidad T", "No aplica", "O(n · 2ⁿ) para construir T"],
            ["Complejidad total", "O(2^(2n-1) · LP(n))", "O(n·2ⁿ + pocas_cand · LP(n))"],
            ["Garantía de optimum", "Sí, siempre", "No garantizada (aproximación)"],
            ["Para n=20", "Congela el PC", "Segundos/minutos"],
            ["Implementación", "bipartition.py, kpartition.py", "geometric.py, GeometricSIA"],
        ]
    )

    warning_box(doc,
        "La solución geométrica puede NO encontrar la bipartición globalmente óptima. "
        "El GeoMIP mismo lo reconoce (Sección 5.2.2): 'dado que el enfoque geométrico "
        "puede no garantizar siempre la bipartición absolutamente óptima, es crucial "
        "cuantificar la calidad de las soluciones encontradas.' "
        "Para el ejemplo ABCD desde estado 100: exhaustivo halla δ=0.0, "
        "geométrico halla δ=0.125 (con 7 candidatos vs 31)."
    )

    presentation_box(doc,
        "Para el jurado: 'El enfoque geométrico es una HEURÍSTICA inteligente. "
        "En lugar de probar todas las divisiones posibles, usa la estructura matemática "
        "de la tabla T para identificar cuáles divisiones son más prometedoras. "
        "Sacrifica la garantía del óptimo global a cambio de una reducción exponencial "
        "en tiempo de cómputo — lo que lo hace viable para sistemas grandes donde "
        "la búsqueda exhaustiva es imposible.'"
    )

    # ─────────────────────────────────────────────────────────────
    # PARTE IX — VERIFICACIÓN Y PRUEBAS
    # ─────────────────────────────────────────────────────────────
    part_banner(doc, "IX", "VERIFICACIÓN Y PRUEBAS")

    # ── 19. Script de verificación ─────────────────────────────
    h1(doc, "19. Script de Verificación Completa")

    body(doc, "Guardar como test_pipeline.py y ejecutar con:")
    code_block(doc, "PYTHONPATH=src /usr/bin/python3.13 test_pipeline.py")

    code_block(doc, '''\
"""test_pipeline.py — Verifica el pipeline completo + k-particiones + geométrico."""
import sys, numpy as np
sys.path.insert(0, 'src')
from tpm_loader import load_tpm, state_to_index, index_to_state
from conditioning import condition_tpm, get_background_state
from marginalization import get_candidate_tpm
from state_node import state_state_to_state_node, state_node_to_state_state
from hypercube import compute_cost_table, hamming_distance
from emd import emd, hamming_distance_matrix
from bipartition import find_optimal_bipartition, build_bipartition_variables, generate_bipartitions
from kpartition import generate_kpartitions, build_partition_variables, count_kpartitions, bell_number, find_optimal_kpartition
from geometric import find_optimal_kpartition_geometric
from controllers.strategies.geometric import GeometricSIA

def ok(cond, msg):
    assert cond, f"FALLO: {msg}"
    print(f"  OK  {msg}")

# ─── 1. Codificación ───
print("=== Codificación ===")
ok(state_to_index([1,0,0]) == 4,  "state_to_index([1,0,0]) = 4")
ok(index_to_state(4, 3) == [1,0,0], "index_to_state(4,3) = [1,0,0]")
ok(state_to_index([]) == 0,       "estado vacío → índice 0")

# ─── 2. TPM ───
print("=== TPM ===")
tpm, n = load_tpm('data/tpm_abcd.csv')
ok(tpm.shape == (16,16), "ABCD: TPM 16×16")
ok(np.allclose(tpm.sum(axis=1), 1.0), "filas suman 1")

# ─── 3-4. Pipeline completo ───
print("=== Condicionamiento + Marginalización ===")
bg_val = get_background_state([1,0,0,0], [3])
cond_tpm, n_c, idx = condition_tpm(tpm, 4, [3], bg_val)
cand_tpm = get_candidate_tpm(cond_tpm, n, idx)
ok(cand_tpm.shape == (8,8), "TPM candidata 8×8")
ok(np.allclose(cand_tpm.sum(axis=1), 1.0), "candidata estocástica")

# ─── 5-6. Estado-nodo + Reconstrucción ───
print("=== Estado-nodo ===")
mats = state_state_to_state_node(cand_tpm, 3)
rec  = state_node_to_state_state(mats, 3)
ok(np.allclose(rec, cand_tpm), "reconstrucción exacta")

# ─── 7. Hamming ───
print("=== Hamming ===")
ok(hamming_distance(0, 7, 3) == 3, "d_H(000,111) = 3")

# ─── 8. EMD ───
print("=== EMD ===")
D = hamming_distance_matrix(3)
p, q = np.zeros(8), np.zeros(8)
p[0]=1.0; q[7]=1.0
ok(abs(emd(p, q, D) - 3.0) < 1e-9, "EMD(000,111) = 3")

# ─── 9. Biparticiones (k=2) ───
print("=== Biparticiones k=2 ===")
bips = list(generate_bipartitions(build_bipartition_variables(3)))
ok(len(bips) == 31, "n=3 → 31 biparticiones (2^(2·3-1)-1)")
res2 = find_optimal_bipartition(cand_tpm, 3, 4, variable_names=['A','B','C'])
ok(res2['n_evaluated'] == 31, "se evaluaron 31")
ok(res2['min_delta'] == 0.0,  "bipartición óptima tiene δ=0")

# ─── 10. K-particiones ───
print("=== K-particiones ===")
ok(count_kpartitions(6, 2) == 31,  "S(6,2) = 31")
ok(count_kpartitions(6, 3) == 90,  "S(6,3) = 90")
ok(bell_number(6) == 203,          "B(6) = 203")
bips3 = list(generate_kpartitions(build_partition_variables(3), 3))
ok(len(bips3) == 90, "n=3, k=3 → 90 particiones")
res3 = find_optimal_kpartition(cand_tpm, 3, 3, 4, variable_names=['A','B','C'])
ok(res3['n_evaluated'] == 90, "se evaluaron 90 particiones (k=3)")
ok(res3['min_delta'] >= 0,    "delta k=3 es no negativo")

# ─── 11. Solución geométrica ───
print("=== GeometricSIA ===")
sia = GeometricSIA(cand_tpm, 3, 4, variable_names=['A','B','C'])
r_geo = sia.aplicar_estrategia(k=2)
ok(r_geo['n_candidates'] > 0, "geométrico identifica candidatos")
ok(r_geo['min_delta'] >= 0,   "delta geométrico no negativo")
costo = sia.calcular_transicion_coste(v=0, i=4, j=0)
ok(costo >= 0, f"T_A[100,000] = {costo:.4f} (no negativo)")

print("\\n✅ Todos los tests pasaron.")''')

    # ── 20. Tabla de archivos ──────────────────────────────────
    h1(doc, "20. Mapa Completo: Módulos, Funciones y Conceptos")

    add_table(doc,
        ["Concepto matemático", "Módulo", "Función / Clase clave"],
        [
            ["Codificación big-endian",          "src/tpm_loader.py",                       "state_to_index, index_to_state"],
            ["Carga y validación de la TPM",     "src/tpm_loader.py",                       "load_tpm"],
            ["Condicionamiento de variables",    "src/conditioning.py",                     "condition_tpm, get_background_state"],
            ["Marginalización columnas (t+1)",   "src/marginalization.py",                  "marginalize_columns, get_candidate_tpm"],
            ["Marginalización filas (t)",        "src/marginalization.py",                  "marginalize_rows"],
            ["Matrices estado-nodo",             "src/state_node.py",                       "state_state_to_state_node"],
            ["Reconstrucción via Kronecker",     "src/state_node.py",                       "state_node_to_state_state"],
            ["Producto tensorial bipartición",   "src/state_node.py",                       "tensor_product_tpm"],
            ["Distancia de Hamming",             "src/hypercube.py",                        "hamming_distance"],
            ["Vecindad del hipercubo",           "src/hypercube.py",                        "get_neighbors, build_hypercube_adjacency"],
            ["Función de costo T (BFS)",         "src/hypercube.py",                        "compute_cost_table, compute_cost_table_from_tpm"],
            ["Earth Mover's Distance",           "src/emd.py",                              "emd, hamming_distance_matrix"],
            ["Biparticiones (k=2) exhaustivo",  "src/bipartition.py",                      "find_optimal_bipartition, generate_bipartitions"],
            ["K-particiones (k≥2) exhaustivo",  "src/kpartition.py",                       "find_optimal_kpartition, generate_kpartitions"],
            ["Números de Stirling y Bell",       "src/kpartition.py",                       "count_kpartitions, bell_number"],
            ["Solución geométrica GeoMIP",       "src/geometric.py",                        "identify_geometric_candidates, find_optimal_kpartition_geometric"],
            ["Clase GeometricSIA (entregable)",  "src/controllers/strategies/geometric.py", "GeometricSIA"],
            ["Interfaz web (Streamlit)",         "app.py",                                  "Pipeline completo con 3 pestañas"],
        ]
    )

    # ── 21. Interpretación de resultados ──────────────────────
    h1(doc, "21. Cómo Interpretar los Resultados")

    add_table(doc,
        ["Resultado", "Significado"],
        [
            ["δ = 0.0", "El sistema es perfectamente separable desde el estado inicial. Las dos partes son completamente independientes."],
            ["δ < 0.1", "Separación casi perfecta. La interdependencia es mínima y probablemente insignificante."],
            ["0.1 ≤ δ ≤ 0.5", "Separación parcial. Existe dependencia causal entre las partes, pero no total."],
            ["δ > 0.5", "Alta interdependencia. No existe una buena bipartición; el sistema está fuertemente acoplado."],
            ["δ = máximo (=n)", "El sistema es completamente no separable: toda la masa debe moverse de 000...0 a 111...1."],
        ]
    )

    question_box(doc,
        "Si δ = 0 para la bipartición óptima, ¿qué significa para el sistema?",
        "Que el sistema puede describirse completamente como dos subsistemas independientes. "
        "La dinámica de S1 no depende de S2 y viceversa desde el estado inicial. "
        "En el contexto de la teoría de la información integrada, significa Φ = 0: "
        "no hay información integrada irreducible en el sistema."
    )

    presentation_box(doc,
        "Para el cierre de la presentación: 'El objetivo final de este proyecto es "
        "calcular eficientemente la separabilidad de sistemas complejos. "
        "Implementamos tres estrategias: búsqueda exhaustiva (garantiza el óptimo "
        "para n pequeño), k-particiones (generaliza a más de 2 grupos), "
        "y la solución geométrica GeoMIP (escalable a sistemas grandes usando "
        "la topología del hipercubo). Juntas cubren el rango completo de aplicaciones.'"
    )

    # ─────────────────────────────────────────────────────────────
    # GUARDAR
    # ─────────────────────────────────────────────────────────────
    out = "BASE_MATEMATICA.docx"
    doc.save(out)
    print(f"Documento guardado: {out}")


if __name__ == "__main__":
    build()
